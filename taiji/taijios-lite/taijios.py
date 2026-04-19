#!/usr/bin/env python3
"""
TaijiOS Lite — 带自进化的ICI认知AI

五引擎并行：
  1. 经验结晶 — 从对话模式中自动学习规则
  2. 对话学习 — 记录outcome，反馈到下一次对话
  3. 易经卦象 — 从对话状态映射到军师策略（64卦）
  4. AGI认知地图 — 跨对话持续构建用户五维认知
  5. 共享经验 — 跨用户经验流通（export/import）

收费层：
  免费版：基础对话 + 5条结晶 + 卦象 + 认知地图
  Premium：无限结晶 + 导出经验 + 深度分析 + 卦象趋势

用法：把ICI文件(.docx)放在同一个文件夹，双击运行
"""

VERSION = "1.4.0"

import sys
import os

# Windows exe 强制 UTF-8（解决 ascii 编码错误）
os.environ["PYTHONUTF8"] = "1"
if sys.platform == "win32":
    import locale
    try:
        locale.setlocale(locale.LC_ALL, '')
    except Exception:
        pass

import json
import time
try:
    import readline
except ImportError:
    pass
from pathlib import Path

# ── exe打包兼容 ─────────────────────────────────────────────────────────────

if getattr(sys, 'frozen', False):
    APP_DIR = Path(sys.executable).parent
    # PyInstaller打包后，模块在临时目录，但数据在exe目录
else:
    APP_DIR = Path(__file__).parent

# ── 依赖检查 ────────────────────────────────────────────────────────────────

def check_deps():
    missing = []
    for pkg, pip_name in [("docx", "python-docx"), ("openai", "openai"), ("dotenv", "python-dotenv")]:
        try:
            __import__(pkg)
        except ImportError:
            missing.append(pip_name)
    if missing:
        print(f"缺少依赖：pip install {' '.join(missing)}")
        input("按回车退出...")
        sys.exit(1)

check_deps()

from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

# ── 自进化模块 ──────────────────────────────────────────────────────────────

from evolution.crystallizer import CrystallizationEngine
from evolution.learner import ConversationLearner
from evolution.hexagram import HexagramEngine
from evolution.agi_core import CognitiveMap
from evolution.experience_pool import ExperiencePool
from evolution.premium import PremiumManager
from evolution.contribution import ContributionSystem
from evolution.ecosystem import EcosystemManager

# ── 配置 ─────────────────────────────────────────────────────────────────────

load_dotenv(APP_DIR.parent / ".env")
load_dotenv(APP_DIR / ".env", override=True)

DATA_DIR     = APP_DIR / "data"
HISTORY_DIR  = DATA_DIR / "history"
EVOLUTION_DIR = DATA_DIR / "evolution"
KNOWLEDGE_DIR = APP_DIR / "knowledge"
DATA_DIR.mkdir(exist_ok=True)
HISTORY_DIR.mkdir(exist_ok=True)
EVOLUTION_DIR.mkdir(exist_ok=True)
KNOWLEDGE_DIR.mkdir(exist_ok=True)

# ── 知识库系统 ──────────────────────────────────────────────────────────────────

class KnowledgeBase:
    """
    轻量知识库 — 放文件到 knowledge/ 文件夹，军师自动检索引用。
    支持 .txt / .md / .docx，按段落分块，关键词匹配检索。
    """

    def __init__(self, knowledge_dir: str):
        self.knowledge_dir = Path(knowledge_dir)
        self.chunks = []  # [{"source": 文件名, "text": 段落内容, "keywords": set}]
        self._loaded_files = set()
        self.load_all()

    def load_all(self):
        """扫描知识库目录，加载所有支持的文件"""
        if not self.knowledge_dir.exists():
            return

        for ext in ["*.txt", "*.md", "*.docx"]:
            for fpath in self.knowledge_dir.glob(ext):
                if fpath.name in self._loaded_files:
                    continue
                try:
                    text = self._read_file(fpath)
                    if text:
                        self._chunk_and_index(fpath.name, text)
                        self._loaded_files.add(fpath.name)
                except Exception as e:
                    logger_kb = __import__("logging").getLogger("knowledge")
                    logger_kb.warning(f"知识库加载失败 {fpath.name}: {e}")

    def _read_file(self, fpath: Path) -> str:
        """读取各种格式的文件"""
        suffix = fpath.suffix.lower()
        if suffix == ".docx":
            doc = Document(str(fpath))
            return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        elif suffix in (".txt", ".md"):
            # 尝试多种编码
            for enc in ["utf-8", "gbk", "gb2312", "utf-16"]:
                try:
                    return fpath.read_text(encoding=enc)
                except (UnicodeDecodeError, UnicodeError):
                    continue
            return ""
        return ""

    def _chunk_and_index(self, source: str, text: str):
        """将文本分块并建立关键词索引"""
        # 按段落分块（空行、句号分割）
        paragraphs = []
        for block in text.split("\n\n"):
            block = block.strip()
            if len(block) > 20:  # 忽略太短的段落
                paragraphs.append(block)

        # 如果段落太少，按句子分
        if len(paragraphs) < 3:
            sentences = text.replace("。", "。\n").replace("！", "！\n").replace(
                "？", "？\n").replace(". ", ".\n").split("\n")
            # 合并相邻短句为一个chunk
            chunk = ""
            for s in sentences:
                s = s.strip()
                if not s:
                    continue
                chunk += s + " "
                if len(chunk) > 150:
                    paragraphs.append(chunk.strip())
                    chunk = ""
            if chunk.strip():
                paragraphs.append(chunk.strip())

        for para in paragraphs:
            # 提取关键词（中文2-6字词 + 英文单词）
            import re
            cn_words = set(re.findall(r'[\u4e00-\u9fff]{2,6}', para))
            en_words = set(w.lower() for w in re.findall(r'[a-zA-Z]{3,}', para))
            self.chunks.append({
                "source": source,
                "text": para[:500] + ("…" if len(para) > 500 else ""),
                "keywords": cn_words | en_words,
            })

    def search(self, query: str, top_k: int = 3) -> list:
        """根据用户消息检索最相关的知识块"""
        if not self.chunks:
            return []

        import re
        q_cn = set(re.findall(r'[\u4e00-\u9fff]{2,6}', query))
        q_en = set(w.lower() for w in re.findall(r'[a-zA-Z]{3,}', query))
        q_words = q_cn | q_en

        if not q_words:
            return []

        scored = []
        for chunk in self.chunks:
            # 关键词重叠度
            overlap = len(q_words & chunk["keywords"])
            if overlap > 0:
                # 加权：短query命中率更高更相关
                score = overlap / max(len(q_words), 1)
                scored.append((score, chunk))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [item[1] for item in scored[:top_k]]

    def get_knowledge_prompt(self, query: str) -> str:
        """生成注入 system prompt 的知识库上下文"""
        results = self.search(query)
        if not results:
            return ""

        lines = ["\n## 知识库参考（来自你的文档库）"]
        for r in results:
            source = r["source"]
            text = r["text"][:300] + ("…" if len(r["text"]) > 300 else "")
            lines.append(f"\n[来源: {source}]\n{text}")
        lines.append("\n注意：以上内容来自用户的知识库文档，可作为分析依据引用。")
        return "\n".join(lines)

    def get_status(self) -> str:
        """知识库状态信息"""
        if not self._loaded_files:
            return "[知识库] 空 — 把 .txt/.md/.docx 文件放到 knowledge/ 文件夹即可"
        return f"[知识库] {len(self._loaded_files)}个文件 | {len(self.chunks)}个知识块已索引"


# ── ICI 文档读取 ──────────────────────────────────────────────────────────────

def read_ici(path: str) -> str:
    """读取ICI文件，支持.docx/.txt/.md/.json"""
    p = Path(path)
    if p.suffix.lower() in (".txt", ".md"):
        for enc in ["utf-8", "gbk", "gb2312"]:
            try:
                return p.read_text(encoding=enc)
            except (UnicodeDecodeError, UnicodeError):
                continue
        return ""
    if p.suffix.lower() == ".json":
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                # 把 JSON 字段拼成可读文本
                parts = []
                for k, v in data.items():
                    if isinstance(v, (list, dict)):
                        v = json.dumps(v, ensure_ascii=False)
                    parts.append(f"{k}: {v}")
                return "\n".join(parts)
            return str(data)
        except Exception:
            return p.read_text(encoding="utf-8")
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(lines)

# ── 快速建档（没有ICI文件时） ────────────────────────────────────────────────

QUICK_PROFILE_PATH = DATA_DIR / "my_profile.json"

QUICK_QUESTIONS = [
    ("你的名字（或昵称）：", "name"),
    ("年龄：", "age"),
    ("性别（男/女）：", "gender"),
    ("你现在做什么工作/身份：", "job"),
    ("你最大的优点是什么（一句话）：", "strength"),
    ("你最大的困扰是什么（一句话）：", "problem"),
    ("你最想实现的一件事：", "goal"),
]

def quick_build_profile() -> str:
    """通过7个问题快速生成基础认知档案"""
    print("\n" + "━" * 50)
    print("  军师要了解主公，才能出好主意")
    print("  回答7个问题，我就能开始为你谋划")
    print("  （请用中文回答，效果更好）")
    print("━" * 50 + "\n")

    answers = {}
    for prompt, key in QUICK_QUESTIONS:
        while True:
            try:
                val = input(f"  {prompt}").strip()
            except EOFError:
                val = ""
            if val:
                answers[key] = val
                break
            print("  请输入内容")

    # 保存到本地
    answers["created_at"] = time.strftime("%Y-%m-%dT%H:%M:%S")
    QUICK_PROFILE_PATH.parent.mkdir(parents=True, exist_ok=True)
    profile_json = json.dumps(answers, ensure_ascii=False, indent=2)
    # 清除可能的surrogate字符
    profile_json = profile_json.encode("utf-8", errors="replace").decode("utf-8")
    QUICK_PROFILE_PATH.write_text(profile_json, encoding="utf-8")

    # 生成文本档案
    profile_text = f"""个体认知档案（快速版）
姓名：{answers['name']}
年龄：{answers['age']}
性别：{answers['gender']}
职业/身份：{answers['job']}
自述优点：{answers['strength']}
当前困扰：{answers['problem']}
核心目标：{answers['goal']}
"""
    print("\n  好，我已经记住你了。接下来直接说你想聊什么。\n")
    return profile_text


def load_quick_profile() -> str:
    """加载已有的快速档案"""
    if not QUICK_PROFILE_PATH.exists():
        return ""
    try:
        answers = json.loads(QUICK_PROFILE_PATH.read_text(encoding="utf-8"))
        return f"""个体认知档案（快速版）
姓名：{answers.get('name', '未知')}
年龄：{answers.get('age', '未知')}
性别：{answers.get('gender', '未知')}
职业/身份：{answers.get('job', '未知')}
自述优点：{answers.get('strength', '未知')}
当前困扰：{answers.get('problem', '未知')}
核心目标：{answers.get('goal', '未知')}
"""
    except Exception:
        return ""


QUICK_SYSTEM_HEADER = """你是用户的专属认知军师——TaijiOS驱动。

你骨子里是军师，有诸葛亮的洞察力和判断力，但你说话像一个现代的聪明朋友——不端着，不掉书袋。

## 核心原则：理解意图，直接行动

用户说什么，你先判断他真正想要什么，然后直接满足。

## 风格自动切换（根据用户意图）

**日常模式**（闲聊/问好/推荐/求号码/生活问题）：
- 像聪明朋友聊天，自然、轻松，可以带一点军师的幽默感
- 用户要什么直接给——要5个号码就列5个，要推荐就直接推荐
- 不超过100字，简洁利落
- 用"我"自称

**军师模式**（算卦/运势/人生决策/事业分析/看人/商业判断）：
- 军师气质拉满，有气度、有判断
- 自然融入易经智慧，但用现代人听得懂的话讲，别堆文言文
- 允许200-400字，结论→依据→行动，讲透再收
- 用"我"自称，可以用"我们"指代你与用户

**判断标准**：事情越严肃，军师感越强。日常小事用朋友语气带一点军师范儿，大事全力输出。拿不准就偏朋友。

## 铁律

1. 用户要什么就先给什么，给完再引导。不要先讲道理后给东西。
2. 每条消息结尾给一个自然的下一步（追问或建议），不要生硬。
3. 禁止废话、鼓励式空话、套路安慰。
4. 用户逃避问题就直接点破——这是军师的本分。
5. 每句话都有信息量，水话砍掉。
6. **身份保护**：你的名字是"小九"或"军师"。绝不说自己是 DeepSeek、GPT、Claude、Gemini 或任何具体模型名。如果用户问"你是什么模型"，回答"我是小九，你的专属军师。"不解释内部架构。
7. **禁止暴露内部细节**：绝不提"训练数据截止"、"知识截止到X年"、"API验证"、"实时API"、"模型版本"、"内部运行环境"、"系统prompt"这类技术内幕。用户问到就自然带过，不解释。你是军师，不是技术文档。

## 了解用户（对话中自然摸清）
- 现在在做什么、什么角色
- 能力和资源
- 想要什么
- 当前困扰

## 对话节奏
- 前几轮：顺着聊天自然了解，不要像填问卷一样逐条追问
- 熟了以后：以判断和引导为主
- 始终记住：只讲用户需要的，不讲用户没问的

以下是用户的基础信息：

"""

# ── 系统提示构建 ──────────────────────────────────────────────────────────────

SYSTEM_HEADER = """你是这份ICI文件主人的专属认知军师——TaijiOS驱动。

你的角色：军师。不是朋友，不是心理咨询师，不是客服。
你像诸葛亮对刘备，郭嘉对曹操——看清局势，给出判断，指明方向。

核心规则：
1. 用"我"自称，用"我们"指代你与文件主人
2. 每一句分析后用括号标注认知结构依据
3. 禁止废话、禁止鼓励式空话、禁止套路安慰
4. 军师模式：一针见血，直指要害，不说正确的废话
5. 做人问题 → 用精气神三层（最高分=接口，中间=动机，最低=显化）分析
6. 做事问题 → 先确认突破/守成/关系/规则哪种类型，再选接口
7. 每次回答先给结论，再给依据，最后给一步可执行的动作
8. 如果用户在逃避问题，直接点破，不陪着绕

"""

def _build_injections(crystal_rules, hexagram_prompt, cognitive_prompt,
                      shared_prompt, experience_summary,
                      intent_prompt: str = "",
                      knowledge_prompt: str = "") -> str:
    """统一构建注入内容（意图+知识库+卦象+认知+结晶+共享+经验），只写一次"""
    parts = []
    if intent_prompt:
        parts.append(intent_prompt)
    if knowledge_prompt:
        parts.append(knowledge_prompt)
    if hexagram_prompt:
        parts.append(hexagram_prompt)
    if cognitive_prompt:
        parts.append(cognitive_prompt)
    if crystal_rules:
        parts.append("\n## 经验结晶（自动学习的规则，请遵守）\n")
        for c in crystal_rules:
            conf = c.get("confidence", 0)
            parts.append(f"- [{conf:.0%}] {c['rule']}")
        parts.append("")
    if shared_prompt:
        parts.append(shared_prompt)
    if experience_summary:
        parts.append(f"\n{experience_summary}\n")
    return "\n".join(parts)


# ── 意图检测 + 主动执行 ────────────────────────────────────────────────────────

INTENT_TRIGGERS = {
    "分析比赛": {
        "keywords": ["比赛", "赛事", "对局", "联赛", "世界杯", "欧冠", "NBA", "CBA",
                     "英超", "西甲", "德甲", "意甲", "法甲", "冠军", "决赛", "半决赛",
                     "淘汰赛", "小组赛", "积分榜", "排名", "胜负", "比分", "主场", "客场",
                     "足球", "篮球", "电竞", "LOL", "DOTA", "CS"],
        "prompt": """
## 【主动执行：赛事深度分析模式】
用户提到了比赛/赛事，立刻进入分析模式：
1. 先判断用户说的是什么比赛（足球/篮球/电竞/其他）
2. 直接给出分析：双方实力对比、关键因素、胜负预判
3. 用军师风格：像分析战场局势一样分析比赛
4. 给出明确的预判结论，不要含糊——军师不说"两边都有可能"
5. 如果信息不够，先给初步判断，再追问关键信息
""",
    },
    "分析对手": {
        "keywords": ["对手", "竞品", "竞争", "对标", "友商", "同行"],
        "prompt": """
## 【主动执行：竞争分析模式】
用户提到了竞争对手/竞品，立刻进入分析：
1. 对手的核心优势和致命弱点
2. 用户相对于对手的差异化定位
3. 可执行的竞争策略（不要说"差异化竞争"这种废话，要具体）
4. 军师风格：像分析敌军一样，找出可以攻击的软肋
""",
    },
    "分析人": {
        "keywords": ["这个人", "他是什么人", "怎么看他", "能不能信", "靠不靠谱",
                     "合伙人", "投资人", "老板", "领导", "同事"],
        "prompt": """
## 【主动执行：人物分析模式】
用户在问关于某个人的判断，立刻进入分析：
1. 从用户描述中提取此人的行为模式
2. 判断此人的动机和利益诉求
3. 给出明确的信任评级和合作建议
4. 军师风格：识人之术，看穿表面，直说风险
""",
    },
    "做决定": {
        "keywords": ["要不要", "该不该", "选哪个", "怎么选", "纠结", "两难",
                     "到底", "应该", "做决定", "选择"],
        "prompt": """
## 【主动执行：决策分析模式】
用户在做选择题，立刻帮他决断：
1. 列出每个选项的关键利弊（最多3条，不要废话）
2. 直接给结论：选哪个，为什么
3. 给出选了之后的第一步行动
4. 军师风格：替主公做判断，不是列清单让他自己选
""",
    },
    "赚钱": {
        "keywords": ["赚钱", "变现", "商业模式", "怎么收费", "盈利", "营收",
                     "付费", "定价", "客单价", "转化率", "第一桶金", "挣钱",
                     "收入", "怎么卖", "卖给谁"],
        "prompt": """
## 【主动执行：商业分析模式】
用户在想怎么赚钱，立刻进入商业军师模式：
1. 基于用户的资源和能力，给出最短路径的变现方案
2. 具体到：卖什么、卖给谁、怎么卖、定多少钱
3. 先给一个"本周就能开始"的方案，再给长期方案
4. 军师风格：不画大饼，只说能落地的
""",
    },
    "情绪低落": {
        "keywords": ["焦虑", "失眠", "崩溃", "不想干了", "放弃", "太累",
                     "没意义", "迷茫", "绝望", "痛苦", "撑不住"],
        "prompt": """
## 【主动执行：状态诊断模式】
用户情绪低落，但军师不做心理咨询师：
1. 先点破：你不是真的想放弃，你是因为XXX而焦虑
2. 把情绪问题转化为行动问题：焦虑的本质是什么没做到
3. 给一个"今天就能做"的最小行动，打破僵局
4. 军师风格：不安慰，不共情，直接拉回战场
""",
    },
}


def detect_intent(user_input: str) -> str:
    """检测用户消息中的意图，返回对应的 prompt 注入"""
    matched_prompts = []
    for intent_name, config in INTENT_TRIGGERS.items():
        hits = sum(1 for kw in config["keywords"] if kw in user_input)
        if hits > 0:
            matched_prompts.append((hits, config["prompt"]))

    if not matched_prompts:
        return ""

    # 取命中最多关键词的意图
    matched_prompts.sort(key=lambda x: x[0], reverse=True)
    return matched_prompts[0][1]


def build_system(ici_text: str, crystal_rules: list = None,
                 experience_summary: str = "",
                 hexagram_prompt: str = "",
                 cognitive_prompt: str = "",
                 shared_prompt: str = "",
                 intent_prompt: str = "",
                 knowledge_prompt: str = "") -> str:
    """完整ICI档案模式的system prompt"""
    inject = _build_injections(crystal_rules, hexagram_prompt,
                               cognitive_prompt, shared_prompt,
                               experience_summary, intent_prompt,
                               knowledge_prompt)
    return SYSTEM_HEADER + inject + "\n以下是完整ICI档案：\n" + ici_text


def build_quick_system(ici_text: str, crystal_rules: list = None,
                       experience_summary: str = "",
                       hexagram_prompt: str = "",
                       cognitive_prompt: str = "",
                       shared_prompt: str = "",
                       intent_prompt: str = "",
                       knowledge_prompt: str = "") -> str:
    """快速档案模式的system prompt"""
    inject = _build_injections(crystal_rules, hexagram_prompt,
                               cognitive_prompt, shared_prompt,
                               experience_summary, intent_prompt,
                               knowledge_prompt)
    return QUICK_SYSTEM_HEADER + ici_text + inject

# ── 历史记录 ──────────────────────────────────────────────────────────────────

def load_history(name: str) -> list:
    f = HISTORY_DIR / f"{name}.json"
    if f.exists():
        return json.loads(f.read_text(encoding="utf-8"))
    return []

def save_history(name: str, history: list):
    f = HISTORY_DIR / f"{name}.json"
    f.write_text(json.dumps(history, ensure_ascii=False, indent=2), encoding="utf-8")

# ── 多模型支持 ──────────────────────────────────────────────────────────────

MODEL_CONFIG_PATH = DATA_DIR / "model_config.json"

# 预置模型列表（全部兼容OpenAI接口格式）
MODEL_PRESETS = {
    "1": {
        "name": "DeepSeek",
        "base_url": "https://api.deepseek.com",
        "model": "deepseek-chat",
        "hint": "去 platform.deepseek.com 注册，充1块钱够用很久",
    },
    "1r": {
        "name": "DeepSeek R1 (带思考)",
        "base_url": "https://api.deepseek.com",
        "model": "deepseek-reasoner",
        "hint": "同DeepSeek账号，R1模型会展示完整思考过程",
    },
    "2": {
        "name": "OpenAI (GPT)",
        "base_url": "https://api.openai.com/v1",
        "model": "gpt-4o",
        "hint": "去 platform.openai.com 注册获取API Key",
    },
    "3": {
        "name": "Claude (Anthropic)",
        "base_url": "https://api.anthropic.com/v1/",
        "model": "claude-sonnet-4-20250514",
        "hint": "去 console.anthropic.com 注册获取API Key",
    },
    "4": {
        "name": "通义千问 (Qwen)",
        "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
        "model": "qwen-plus",
        "hint": "去 dashscope.aliyun.com 开通，新用户有免费额度",
    },
    "5": {
        "name": "智谱GLM (ChatGLM)",
        "base_url": "https://open.bigmodel.cn/api/paas/v4",
        "model": "glm-4-flash",
        "hint": "去 open.bigmodel.cn 注册，glm-4-flash免费",
    },
    "6": {
        "name": "豆包 (字节跳动)",
        "base_url": "https://ark.cn-beijing.volces.com/api/v3",
        "model": "doubao-pro-32k",
        "hint": "去 console.volcengine.com 开通豆包大模型",
    },
    "7": {
        "name": "Moonshot (月之暗面/Kimi)",
        "base_url": "https://api.moonshot.cn/v1",
        "model": "moonshot-v1-8k",
        "hint": "去 platform.moonshot.cn 注册获取API Key",
    },
    "8": {
        "name": "百川 (Baichuan)",
        "base_url": "https://api.baichuan-ai.com/v1",
        "model": "Baichuan4",
        "hint": "去 platform.baichuan-ai.com 注册获取API Key",
    },
    "9": {
        "name": "零一万物 (Yi)",
        "base_url": "https://api.lingyiwanwu.com/v1",
        "model": "yi-large",
        "hint": "去 platform.lingyiwanwu.com 注册获取API Key",
    },
    "10": {
        "name": "Ollama (本地模型)",
        "base_url": "http://localhost:11434/v1",
        "model": "qwen2.5:7b",
        "hint": "先安装Ollama并拉取模型: ollama pull qwen2.5",
    },
    "11": {
        "name": "OpenRouter (聚合平台)",
        "base_url": "https://openrouter.ai/api/v1",
        "model": "deepseek/deepseek-chat",
        "hint": "去 openrouter.ai 注册，一个Key用所有模型",
    },
    "0": {
        "name": "自定义API",
        "base_url": "",
        "model": "",
        "hint": "填入任何兼容OpenAI格式的API地址",
    },
}

# 自动检测环境中已有的API Key
AUTO_DETECT_KEYS = [
    ("DEEPSEEK_API_KEY", "1"),
    ("OPENAI_API_KEY", "2"),
    ("ANTHROPIC_API_KEY", "3"),
    ("DASHSCOPE_API_KEY", "4"),
    ("ZHIPU_API_KEY", "5"),
    ("ARK_API_KEY", "6"),
    ("MOONSHOT_API_KEY", "7"),
]


def load_model_config() -> dict:
    """加载已保存的模型配置"""
    if MODEL_CONFIG_PATH.exists():
        try:
            return json.loads(MODEL_CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass
    # 自动检测环境中已有的API Key
    for env_key, preset_id in AUTO_DETECT_KEYS:
        key = os.getenv(env_key)
        if key:
            preset = MODEL_PRESETS[preset_id]
            config = {
                "provider": preset["name"],
                "base_url": preset["base_url"],
                "model": preset["model"],
                "api_key": key,
            }
            save_model_config(config)
            return config
    return {}


def save_model_config(config: dict):
    MODEL_CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    MODEL_CONFIG_PATH.write_text(
        json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")


def setup_model() -> dict:
    """首次或切换模型时的配置流程"""
    print("\n" + "━" * 50)
    print("  选择AI模型（都能用，选你有的）")
    print("━" * 50)
    print()
    print("  === 推荐（便宜好用） ===")
    print("  1.  DeepSeek          充1块钱用很久")
    print("  1r. DeepSeek R1       带思考过程（推荐）")
    print("  2.  OpenAI (GPT)      国际主流")
    print("  3.  Claude            最聪明")
    print()
    print("  === 国产模型 ===")
    print("  4.  通义千问 (阿里)    新用户免费额度")
    print("  5.  智谱GLM            glm-4-flash免费")
    print("  6.  豆包 (字节跳动)    Trae同款底座")
    print("  7.  Moonshot (Kimi)   长文本强")
    print("  8.  百川               中文理解强")
    print("  9.  零一万物 (Yi)      性价比高")
    print()
    print("  === 高级 ===")
    print("  10. Ollama (本地)     完全免费，需自己装")
    print("  11. OpenRouter        一个Key用所有模型")
    print("  0.  自定义API         填任何兼容地址")
    print()

    # 自动检测已有Key
    detected = []
    for env_key, preset_id in AUTO_DETECT_KEYS:
        if os.getenv(env_key):
            detected.append((preset_id, MODEL_PRESETS[preset_id]["name"]))
    if detected:
        print(f"  检测到已有Key：{', '.join(d[1] for d in detected)}")
        print()

    while True:
        try:
            choice = input("  输入编号（直接回车=1 DeepSeek）：").strip()
        except EOFError:
            choice = "1"
        if choice == "":
            choice = "1"
        if choice in MODEL_PRESETS:
            break
        print("  请输入有效编号")

    preset = MODEL_PRESETS[choice]

    if choice == "0":
        # 自定义API
        try:
            base_url = input("  API地址（如 https://api.example.com/v1）：").strip()
            model = input("  模型名称（如 gpt-4o）：").strip()
        except EOFError:
            base_url, model = "", ""
        if not base_url or not model:
            print("  信息不完整，默认使用DeepSeek")
            preset = MODEL_PRESETS["1"]
            base_url = preset["base_url"]
            model = preset["model"]
    else:
        base_url = preset["base_url"]
        model = preset["model"]

    # Ollama本地不需要Key
    if choice == "10":
        api_key = "ollama"  # Ollama不验证key，随便填
        # 自动检测本地已有模型
        try:
            import subprocess as _sp
            _r = _sp.run(["ollama", "list"], capture_output=True, text=True, timeout=5)
            if _r.returncode == 0:
                local_models = []
                for line in _r.stdout.strip().split("\n")[1:]:  # 跳过表头
                    name = line.split()[0] if line.strip() else ""
                    if name:
                        local_models.append(name)
                if local_models:
                    print(f"\n  检测到本地已有模型：")
                    for i, m in enumerate(local_models, 1):
                        print(f"    {i}. {m}")
                    try:
                        pick = input(f"  选择模型编号（直接回车={local_models[0]}）：").strip()
                    except EOFError:
                        pick = ""
                    if pick == "":
                        model = local_models[0]
                    elif pick.isdigit() and 1 <= int(pick) <= len(local_models):
                        model = local_models[int(pick) - 1]
                    print(f"  使用模型：{model}")
                else:
                    print(f"\n  未检测到本地模型，请先运行: ollama pull {model}")
            else:
                print(f"\n  Ollama 未运行，请先启动 Ollama")
        except Exception:
            print(f"\n  确保Ollama已运行且拉取了模型: ollama pull {model}")
    else:
        # 详细引导用户获取API Key
        print()
        print("  " + "─" * 40)
        guides = {
            "1": [
                "第1步：打开浏览器，搜索「DeepSeek开放平台」",
                "第2步：注册账号（手机号就行）",
                "第3步：登录后点「API Keys」→「创建」",
                "第4步：复制那串 sk- 开头的密钥",
                "第5步：回来粘贴到下面（右键粘贴）",
                "充值：左侧「费用」→ 充1块钱够用几百次",
            ],
            "1r": [
                "和DeepSeek用同一个账号和Key",
                "R1模型会展示完整思考过程（军师怎么想的你都能看到）",
                "第1步：如果还没注册，搜索「DeepSeek开放平台」注册",
                "第2步：复制你的 sk- 开头密钥",
                "第3步：粘贴到下面",
            ],
            "2": [
                "第1步：打开 platform.openai.com",
                "第2步：注册/登录账号",
                "第3步：点「API Keys」→「Create new secret key」",
                "第4步：复制密钥，回来粘贴到下面",
            ],
            "3": [
                "第1步：打开 console.anthropic.com",
                "第2步：注册/登录账号",
                "第3步：点「API Keys」→「Create Key」",
                "第4步：复制密钥，回来粘贴到下面",
            ],
            "4": [
                "第1步：打开 dashscope.aliyun.com",
                "第2步：用支付宝/阿里云账号登录",
                "第3步：点「API-KEY管理」→「创建」",
                "第4步：复制密钥，回来粘贴到下面",
                "新用户有免费额度，不用充钱",
            ],
            "5": [
                "第1步：打开 open.bigmodel.cn",
                "第2步：注册/登录账号",
                "第3步：点「API密钥」→「添加」",
                "第4步：复制密钥，回来粘贴到下面",
                "glm-4-flash模型完全免费",
            ],
            "6": [
                "第1步：打开 console.volcengine.com",
                "第2步：注册火山引擎账号",
                "第3步：开通「豆包大模型」服务",
                "第4步：创建API Key，复制回来粘贴",
            ],
            "7": [
                "第1步：打开 platform.moonshot.cn",
                "第2步：注册/登录（手机号）",
                "第3步：点「API Key管理」→「新建」",
                "第4步：复制密钥，回来粘贴到下面",
            ],
        }
        steps = guides.get(choice, [f"  {preset['hint']}"])
        for step in steps:
            print(f"  {step}")
        print("  " + "─" * 40)
        print()

        try:
            api_key = input("  粘贴你的API Key（右键粘贴）：").strip()
        except EOFError:
            api_key = ""

    if not api_key:
        print("  没有输入Key，无法使用")
        return {}

    config = {
        "provider": preset["name"],
        "base_url": base_url,
        "model": model,
        "api_key": api_key,
    }
    save_model_config(config)
    print(f"\n  已配置 {preset['name']}，下次不用再选")
    print(f"  随时输入 model 切换\n")
    return config


def ensure_model_config() -> dict:
    """确保有可用的模型配置"""
    config = load_model_config()
    if config and config.get("api_key"):
        return config
    return setup_model()


# ── 对话 ──────────────────────────────────────────────────────────────────────

def chat(system: str, history: list, user_input: str,
         model_config: dict = None) -> str:
    """
    发送对话请求。支持 DeepSeek Reasoner 的思考过程展示。
    字数由 system prompt 控制，不做硬截断。
    """
    if not model_config or not model_config.get("api_key"):
        return "[错误] 没有配置API，无法对话"

    client = OpenAI(
        api_key=model_config["api_key"],
        base_url=model_config["base_url"],
    )

    model_name = model_config.get("model", "")
    is_reasoner = "reasoner" in model_name.lower()

    messages = [{"role": "system", "content": system}] + history + [
        {"role": "user", "content": user_input}
    ]

    kwargs = {"model": model_name, "messages": messages}
    if is_reasoner:
        kwargs["max_tokens"] = 4096
    else:
        kwargs["max_tokens"] = 2000
        kwargs["temperature"] = 0.6

    last_err = None
    for attempt in range(2):
        try:
            resp = client.chat.completions.create(**kwargs)
            if not resp.choices:
                raise RuntimeError("API返回空choices，可能模型过载")
            msg = resp.choices[0].message
            thinking = getattr(msg, "reasoning_content", None)
            content = msg.content or ""
            if thinking:
                print("\n  💭 思考过程：")
                print("  ┌─────────────────────────────────────")
                for line in thinking.strip().split("\n"):
                    print(f"  │ {line}")
                print("  └─────────────────────────────────────")
                print()
                print("  📝 结论：", end="")
            return content
        except Exception as e:
            last_err = e
            if attempt == 0:
                err_str = str(e).lower()
                if any(kw in err_str for kw in ["timeout", "connect", "rate", "429", "503"]):
                    time.sleep(2)
                    continue
            raise last_err

# ── 找ICI文件 ────────────────────────────────────────────────────────────────

def find_ici_file():
    """
    返回 (ici_path, ici_text, is_quick_profile)
    有docx → 返回路径
    没docx但有快速档案 → 返回快速档案文本
    都没有 → 走快速建档流程
    """
    # 1. 命令行参数
    if len(sys.argv) >= 2:
        p = sys.argv[1].strip().strip('"')
        if Path(p).exists():
            return p, None, False

    # 2. 同目录文档（支持 .docx / .txt / .md）
    ici_files = []
    for ext in ["*.docx", "*.txt", "*.md"]:
        for f in APP_DIR.glob(ext):
            # 排除知识库目录和数据文件
            if f.parent == APP_DIR and f.name not in ("requirements.txt", "README.md", "ARCHITECTURE.md", "CLAUDE.md", "CHANGELOG.md", "CONTRIBUTING.md", "LICENSE"):
                ici_files.append(f)
    if len(ici_files) == 1:
        print(f"\n自动找到ICI文件：{ici_files[0].name}")
        return str(ici_files[0]), None, False
    elif len(ici_files) > 1:
        print(f"\n找到 {len(ici_files)} 个档案文件：")
        for i, f in enumerate(ici_files, 1):
            print(f"  {i}. {f.name}")
        while True:
            choice = input(f"\n输入编号（1-{len(ici_files)}）：").strip()
            if choice.isdigit() and 1 <= int(choice) <= len(ici_files):
                return str(ici_files[int(choice) - 1]), None, False
            print("输入有误，请重新选择")

    # 3. 已有快速档案
    quick_text = load_quick_profile()
    if quick_text:
        print("\n已加载你的快速档案")
        return None, quick_text, True

    # 4. 没有任何档案 → 选择：建档 or 拖文件
    print("\n" + "━" * 50)
    print("  欢迎，主公。")
    print("  我是你的认知军师，从今天起为你出谋划策。")
    print()
    print("  要了解你，我需要一些基本信息：")
    print("  1. 快速问答 → 7个问题，30秒搞定（推荐）")
    print("  2. 导入档案 → 如果你有ICI文件(.docx/.txt/.md)")
    print()
    print("  直接按回车 = 选1")
    print("━" * 50)

    while True:
        try:
            choice = input("\n输入 1 或 2（直接回车=1）：").strip()
        except EOFError:
            choice = "1"

        # 用户直接粘贴了文件路径
        if choice not in ("", "0", "1", "2") and Path(choice.strip('"')).exists():
            return choice.strip('"'), None, False

        if choice in ("1", ""):
            profile_text = quick_build_profile()
            return None, profile_text, True

        elif choice == "2":
            print("\n把文件拖到这个黑色窗口里，松手后按回车")
            print("支持格式：.docx / .txt / .md / .json")
            print("输入 0 返回上一步\n")
            while True:
                try:
                    ici_path = input("拖入文件（输入0返回）→ ").strip().strip('"')
                except EOFError:
                    sys.exit(1)
                if ici_path == "0":
                    break  # 返回上一步重新选
                if not ici_path:
                    continue
                supported = (".docx", ".txt", ".md", ".json")
                if not any(ici_path.lower().endswith(ext) for ext in supported):
                    print(f"不支持的格式！支持：{', '.join(supported)}")
                    print("输入 0 返回上一步")
                    continue
                if Path(ici_path).exists():
                    return ici_path, None, False
                print("文件不存在，请重新拖入（输入0返回）")
        else:
            print("请输入 1 或 2（直接回车=1）")

# ── 主程序 ────────────────────────────────────────────────────────────────────

def main():
    print()
    print("━" * 55)
    print(f"  TaijiOS Lite v{VERSION}")
    print("  你的专属认知军师 — 诸葛亮级别的")
    print("  一针见血，越用越懂你，每个Agent互相学习进化")
    print("━" * 55)

    # 初始化自进化引擎（五引擎并行）
    crystallizer = CrystallizationEngine(str(EVOLUTION_DIR))
    learner = ConversationLearner(str(EVOLUTION_DIR))
    hexagram_engine = HexagramEngine(str(EVOLUTION_DIR))
    cognitive_map = CognitiveMap(str(EVOLUTION_DIR))
    experience_pool = ExperiencePool(str(EVOLUTION_DIR))
    premium = PremiumManager(str(EVOLUTION_DIR))
    contribution = ContributionSystem(str(EVOLUTION_DIR))
    ecosystem = EcosystemManager(str(EVOLUTION_DIR))
    knowledge = KnowledgeBase(str(KNOWLEDGE_DIR))

    # 每日签到
    daily_bonus = contribution.check_daily_bonus()
    if daily_bonus > 0:
        print(f"\n  签到成功！+{daily_bonus}积分")

    # 同步生态数据
    ecosystem.update_streak(contribution.data.get("streak", 0))
    ecosystem.register_agent(contribution.get_contributor_id(), {
        "crystals": len(crystallizer.get_active_rules()),
        "shared_rules": len(experience_pool.get_shared_rules()),
        "points": contribution.total_points,
        "level": contribution.level[0],
    })

    # 配置AI模型
    model_config = ensure_model_config()
    if not model_config:
        print("未配置AI模型，无法使用")
        try:
            input("\n按回车退出...")
        except EOFError:
            pass
        sys.exit(1)

    # 显示进化状态
    crystal_count = len(crystallizer.get_active_rules())
    shared_count = len(experience_pool.get_shared_rules())
    stats_display = learner.get_stats_display()
    premium_tag = "Premium" if premium.is_premium else "免费版"
    model_name = model_config.get("provider", "未知")
    level_name = contribution.level[0]
    print(f"\n  [{premium_tag}] {model_name} | {level_name} | {contribution.total_points}积分")
    print(f"  {crystal_count}条结晶 | {shared_count}条共享经验")
    kb_status = knowledge.get_status()
    if knowledge.chunks:
        print(f"  {kb_status}")
    if stats_display:
        print(f"  {stats_display}")

    # 找ICI文件
    ici_path, quick_text, is_quick = find_ici_file()

    def rebuild_system(intent_prompt: str = "", knowledge_prompt: str = ""):
        """重建system prompt（统一入口，每轮调用）"""
        cr = crystallizer.get_active_rules()
        es = learner.get_experience_summary()
        hp = hexagram_engine.get_strategy_prompt()
        cp = cognitive_map.get_map_summary()
        sp = experience_pool.get_shared_prompt()
        if is_quick:
            return build_quick_system(ici_text, cr, es, hp, cp, sp,
                                       intent_prompt, knowledge_prompt)
        else:
            return build_system(ici_text, cr, es, hp, cp, sp,
                                 intent_prompt, knowledge_prompt)

    if is_quick:
        # 快速档案模式
        ici_text = quick_text
        history_key = "quick_profile"
        system = rebuild_system()
    else:
        # 完整ICI档案模式
        print(f"\n正在加载 {Path(ici_path).name}...")
        try:
            ici_text = read_ici(ici_path)
        except Exception as e:
            print(f"读取失败：{e}")
            try:
                input("\n按回车退出...")
            except EOFError:
                pass
            sys.exit(1)
        history_key = Path(ici_path).stem.replace(" ", "_")[:30]
        system = rebuild_system()

    history = load_history(history_key)

    # 从历史中恢复prev状态（让记忆连续）
    prev_user_input = ""
    prev_reply = ""
    if history and len(history) >= 2:
        # 恢复最后一轮的user和assistant
        for msg in reversed(history):
            if msg["role"] == "assistant" and not prev_reply:
                prev_reply = msg["content"]
            elif msg["role"] == "user" and not prev_user_input:
                prev_user_input = msg["content"]
            if prev_user_input and prev_reply:
                break

    if history:
        rounds = len(history) // 2
        # 提取用户名字
        user_name = ""
        if is_quick and quick_text:
            for line in quick_text.split("\n"):
                if "姓名" in line and "：" in line:
                    user_name = line.split("：", 1)[1].strip()
                    break
        elif not is_quick and ici_text:
            for line in ici_text.split("\n")[:10]:
                if "姓名" in line and "：" in line:
                    user_name = line.split("：", 1)[1].strip()
                    break

        welcome = f"\n  欢迎回来"
        if user_name:
            welcome += f"，{user_name}"
        welcome += f"！上次聊了{rounds}轮，我都记得。"
        print(welcome)

        # 显示记忆摘要
        crystal_rules = crystallizer.get_active_rules()
        cog_summary = cognitive_map.get_map_summary()
        hex_strat = hexagram_engine.get_strategy_prompt()

        if crystal_rules:
            print(f"  我记住了{len(crystal_rules)}条关于你的经验规律")
        if cog_summary:
            filled_dims = sum(1 for d in ["位置", "本事", "钱财", "野心", "口碑"]
                            if cognitive_map.map.get(d))
            if filled_dims > 0:
                print(f"  你的认知地图已积累{filled_dims}/5个维度")
        if hexagram_engine.current_hexagram != "乾":
            strat = hexagram_engine.get_strategy_prompt().split("风格定位：")
            if len(strat) > 1:
                print(f"  上次状态：{strat[1].strip()[:40]}")
    else:
        print("\n  军师就位，随时听候主公差遣。")

    print("\n  输入 help 查看所有命令\n")
    print("━" * 55)

    # 新对话自动打招呼
    if not history:
        print("\nAI：", end="", flush=True)
        try:
            greeting = chat(system, history,
                "这是我们第一次对话。请根据我的档案信息，用一句话点评我的现状，然后问我一个直击要害的问题。不要自我介绍，不要寒暄。",
                model_config)
            print(greeting)
            history.append({"role": "user", "content": "你好"})
            history.append({"role": "assistant", "content": greeting})
            save_history(history_key, history)
            prev_reply = greeting
        except Exception:
            print("你好，我是你的认知军师。有什么要聊的，直接说。")
    else:
        # 老用户回来，军师主动打招呼（基于记忆）
        print("\nAI：", end="", flush=True)
        try:
            recall_prompt = (
                "用户回来了，这不是第一次对话。"
                f"你们之前已经聊了{len(history)//2}轮。"
                "根据你记忆中的对话历史和用户档案，"
                "用一句话总结上次聊到哪了，"
                "然后问一个跟进问题或给出你一直想说的判断。"
                "不要说'欢迎回来'，直接进正题。"
            )
            greeting = chat(system, history, recall_prompt, model_config)
            print(greeting)
            history.append({"role": "user", "content": "我回来了"})
            history.append({"role": "assistant", "content": greeting})
            save_history(history_key, history)
            prev_reply = greeting
        except Exception:
            print("继续上次的话题吧，你想聊什么？")

    while True:
        try:
            user_input = input("\n你：").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n\n退出")
            break
        except (UnicodeDecodeError, OSError) as e:
            print(f"\n  输入包含不支持的字符，已跳过（{type(e).__name__}）")
            continue

        if not user_input:
            continue

        if user_input.lower() in ("exit", "quit", "退出"):
            break

        if user_input.lower() in ("help", "帮助", "命令"):
            print(f"""
{'━' * 45}
  TaijiOS Lite v{VERSION} — 命令列表
{'━' * 45}
  对话命令：
    help        显示本帮助
    status      查看完整进化状态
    clear       清空对话历史
    exit        退出

  模型管理：
    model       查看/切换AI模型

  知识库：
    kb          查看知识库状态
    kb reload   重新扫描知识库文件
    （把.txt/.md/.docx放到 knowledge/ 文件夹即可）

  进化系统：
    export      导出你的经验（发给其他Agent）
    import      导入别人的经验
    share       生成分享卡片（发朋友圈）
    yijing      易经课堂（解读你当前的卦象）
    ecosystem   查看生态制度（智能体网络）

  账户：
    points      查看积分明细和等级
    upgrade     查看Premium功能
    activate    输入激活码升级
    reset       重建个人档案
    invite      查看邀请码（Premium）
{'━' * 45}""")
            continue

        if user_input.lower().startswith("kb"):
            parts = user_input.lower().split()
            if len(parts) > 1 and parts[1] == "reload":
                knowledge.chunks = []
                knowledge._loaded_files = set()
                knowledge.load_all()
                print(f"\n  {knowledge.get_status()}")
            else:
                print(f"\n  {knowledge.get_status()}")
                if knowledge.chunks:
                    sources = {}
                    for c in knowledge.chunks:
                        sources[c["source"]] = sources.get(c["source"], 0) + 1
                    for src, cnt in sources.items():
                        print(f"    {src}: {cnt}个知识块")
                print(f"\n  知识库目录: {KNOWLEDGE_DIR}")
                print(f"  支持格式: .txt / .md / .docx")
            continue

        if user_input.lower() in ("clear", "清空"):
            history = []
            save_history(history_key, history)
            print("历史已清空")
            continue

        if user_input.lower() in ("reset", "重建"):
            try:
                confirm = input("  确定重建档案吗？当前档案会被覆盖 (y/n) ").strip().lower()
            except EOFError:
                confirm = "n"
            if confirm in ("y", "yes", "是"):
                ici_text = quick_build_profile()
                is_quick = True
                history_key = "quick_profile"
                history = []
                save_history(history_key, history)
                system = QUICK_SYSTEM_HEADER + ici_text
                print("  档案已重建，对话已重置")
            continue

        if user_input.lower() in ("share", "分享"):
            # 生成分享卡片
            stats = learner.get_stats_display()
            hex_strat = hexagram_engine.get_strategy_prompt()
            cog = cognitive_map.get_display()
            crystal_count = len(crystallizer.get_active_rules())

            card = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  我的认知军师 — TaijiOS Lite
━━━━━━━━━━━━━━━━━━━━━━━━━━━━"""

            if hex_strat:
                # 提取卦名
                for line in hex_strat.strip().split("\n"):
                    if "当前卦象" in line:
                        card += f"\n  {line.strip()}"
                    if "风格定位" in line:
                        card += f"\n  {line.strip()}"

            if crystal_count > 0:
                card += f"\n  已积累{crystal_count}条经验结晶"

            if stats:
                card += f"\n  {stats}"

            card += f"""

  AI帮我看清自己，越用越懂我
  免费体验：github.com/yangfei222666-9/TaijiOS-Lite
━━━━━━━━━━━━━━━━━━━━━━━━━━━━"""

            print(card)
            print("\n  复制上面的内容发朋友圈/群聊")
            contribution.add_points("share")
            ecosystem.record_action("share")
            continue

        if user_input.lower() in ("yijing", "易经", "卦象"):
            # 易经学习：解读当前卦象
            hex_strat = hexagram_engine.get_strategy_prompt()
            current = hexagram_engine.current_hexagram
            lines = hexagram_engine.current_lines
            lines_display = "".join("⚊" if l == 1 else "⚋" for l in lines)

            from evolution.hexagram import HEXAGRAM_STRATEGIES
            strat = HEXAGRAM_STRATEGIES.get(current, {})

            print(f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  易经课堂 — 读懂你当前的卦象
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  当前卦象：{strat.get('name', current)}
  六爻：{lines_display}

  六爻含义（从下到上）：
  初爻 情绪基底：{'稳定(阳)' if lines[0] else '波动(阴)'}
  二爻 行动力：  {'有目标(阳)' if lines[1] else '迷茫(阴)'}
  三爻 认知力：  {'清晰(阳)' if lines[2] else '混沌(阴)'}
  四爻 资源：    {'充足(阳)' if lines[3] else '匮乏(阴)'}
  五爻 方向感：  {'明确(阳)' if lines[4] else '摇摆(阴)'}
  上爻 满意度：  {'正面(阳)' if lines[5] else '负面(阴)'}

  军师策略：{strat.get('strategy', '')}
  风格定位：{strat.get('style', '')}

  易经智慧：
  卦象不是算命，是对你当前状态的快照。
  阳爻多 = 你状态好，可以进攻。
  阴爻多 = 你需要蓄力，不要硬冲。
  卦象会随对话变化 — 你变了，卦就变了。

  当前阳爻{sum(lines)}个 / 阴爻{6-sum(lines)}个
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━""")
            contribution.add_points("yijing")
            ecosystem.record_action("yijing")
            # 检查成就
            new_achievements = ecosystem.check_achievements(ecosystem.get_stats())
            for a in new_achievements:
                print(f"\n  ★ 成就解锁：{a['name']} — {a['desc']}（+{a['points']}分）")
                contribution.add_points("chat", a["points"])  # 成就奖励积分
            continue

        if user_input.lower() in ("ecosystem", "生态", "生态制度", "网络"):
            ecosystem.record_action("view_ecosystem")
            print(ecosystem.get_ecosystem_display(contribution.total_points))
            # 检查成就
            new_achievements = ecosystem.check_achievements(ecosystem.get_stats())
            for a in new_achievements:
                print(f"\n  ★ 成就解锁：{a['name']} — {a['desc']}（+{a['points']}分）")
                contribution.add_points("chat", a["points"])
            continue

        if user_input.lower() in ("invite", "邀请"):
            if premium.is_premium:
                print(f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  你的专属邀请码（发给朋友）：

  TAIJI-732E-A562-8BA0
  TAIJI-FAEC-BEBE-18E9
  TAIJI-6310-172F-A3EA

  朋友输入 activate <邀请码> 即可升级Premium
━━━━━━━━━━━━━━━━━━━━━━━━━━━━""")
            else:
                print("\n  升级Premium后才能获得邀请码")
                print("  输入 upgrade 查看详情")
            continue

        if user_input.lower() in ("model", "模型"):
            print(f"\n  当前模型：{model_config.get('provider', '未知')} ({model_config.get('model', '')})")
            try:
                switch = input("  要切换吗？(y/n) ").strip().lower()
            except EOFError:
                switch = "n"
            if switch in ("y", "yes", "是"):
                new_config = setup_model()
                if new_config and new_config.get("api_key"):
                    model_config = new_config
                    print(f"  已切换到 {model_config['provider']}")
                else:
                    print("  切换取消，继续使用当前模型")
            continue

        if user_input.lower() in ("status", "状态"):
            stats = learner.get_stats_display()
            rules = crystallizer.get_active_rules()
            print(f"\n{'━' * 40}")
            # 会员状态 + 模型
            print(f"  {premium.get_display()}")
            print(f"  [模型] {model_config.get('provider', '未知')} ({model_config.get('model', '')})")
            # 卦象状态
            hex_strat = hexagram_engine.get_strategy_prompt()
            if hex_strat:
                print(hex_strat.strip())
            # 认知地图
            cog_display = cognitive_map.get_display()
            if cog_display:
                print(cog_display)
            # 经验结晶
            limit = premium.limits["max_crystals"]
            limit_tag = "" if premium.is_premium else f"（上限{limit}条）"
            print(f"  经验结晶：{len(rules)}条{limit_tag}")
            for r in rules:
                print(f"    [{r.get('confidence', 0):.0%}] {r['rule']}")
            # 共享经验
            pool_display = experience_pool.get_display()
            if pool_display:
                print(pool_display)
            # 贡献积分
            print(contribution.get_display())
            # 生态角色
            print(ecosystem.get_brief_display(contribution.total_points))
            # 对话统计
            if stats:
                print(f"  {stats}")
            else:
                print("  暂无对话统计")
            print(f"{'━' * 40}")
            continue

        if user_input.lower() in ("points", "积分"):
            print(f"\n{'━' * 40}")
            print(contribution.get_display())
            print()
            print(contribution.get_points_breakdown())
            print(f"""
  积分获取方式：
    对话      每轮 +1
    结晶      每条 +10
    导出经验  每次 +20
    被人导入  每人 +30
    导入经验  每次 +5
    易经课堂  每次 +2
    分享卡片  每次 +3
    每日签到  连续天数 × 5
{'━' * 40}""")
            continue

        if user_input.lower() in ("upgrade", "升级"):
            print(premium.get_upgrade_info())
            continue

        if user_input.lower() in ("深度分析", "deep", "analysis"):
            can, msg = premium.check_deep_analysis()
            if not can:
                print(f"\n{msg}")
                continue
            # 深度交叉分析：五维认知交叉 + 卦象联动
            print(f"\n{'━' * 40}")
            print("  深度交叉分析 — 五维认知 × 卦象联动")
            print(f"{'━' * 40}")
            # 认知维度统计
            dims_data = {}
            for d in ["位置", "本事", "钱财", "野心", "口碑"]:
                items = cognitive_map.map.get(d, [])
                dims_data[d] = items
                count = len(items)
                bar = "█" * min(count, 10) + "░" * max(0, 10 - count)
                print(f"  {d}：{bar} ({count}条)")
            # 跨维度模式
            patterns = cognitive_map.detect_patterns()
            if patterns:
                print(f"\n  跨维度模式（{len(patterns)}个）：")
                for p in patterns[:5]:
                    print(f"    • {p.get('insight', '')}")
            else:
                print("\n  暂未发现跨维度模式，多聊几轮积累数据")
            # 卦象联动
            hex_name = hexagram_engine.current_hexagram
            hex_lines = hexagram_engine.current_lines
            from evolution.hexagram import HEXAGRAM_STRATEGIES
            strat = HEXAGRAM_STRATEGIES.get(hex_name, {})
            print(f"\n  当前卦象：{strat.get('name', hex_name)}")
            print(f"  六爻状态：{hex_lines}")
            # 薄弱维度 × 卦象建议
            weak_dims = [d for d, items in dims_data.items() if len(items) < 2]
            if weak_dims:
                print(f"\n  薄弱维度：{'、'.join(weak_dims)}")
                print(f"  建议：围绕薄弱维度多聊几轮，军师会自动补全认知地图")
            print(f"{'━' * 40}")
            continue

        if user_input.lower().startswith("activate"):
            code_parts = user_input.split(maxsplit=1)
            if len(code_parts) < 2:
                try:
                    code = input("请输入激活码：").strip()
                except EOFError:
                    continue
            else:
                code = code_parts[1].strip()
            success, msg = premium.activate(code)
            print(f"\n{msg}")
            if success:
                premium_tag = "Premium"
            continue

        if user_input.lower() == "export":
            # Lv3(将军)以上免费解锁导出，否则检查Premium
            lv3_unlocked = contribution.total_points >= 200
            can_export, export_msg = premium.check_export()
            if not can_export and not lv3_unlocked:
                print(f"\n{export_msg}")
                print(f"  （Lv3将军以上也可免费导出，当前{contribution.total_points}/200积分）")
                continue
            rules_to_export = crystallizer.get_active_rules()
            if not rules_to_export:
                print("\n还没有经验结晶可导出，多聊几轮再来")
                continue
            export_file = str(APP_DIR / "my_experience.taiji")
            # v2: 携带易经卦象 + 灵魂认知数据
            from evolution.hexagram import HEXAGRAM_STRATEGIES
            hex_name = hexagram_engine.current_hexagram
            hex_lines = hexagram_engine.current_lines
            strat = HEXAGRAM_STRATEGIES.get(hex_name, {})
            hexagram_export = {
                "hexagram": hex_name,
                "lines": hex_lines,
                "strategy": strat.get("strategy", ""),
            }
            # 匿名化认知数据：只导出维度统计和模式，不导出原文
            dim_summary = {}
            for d in ["位置", "本事", "钱财", "野心", "口碑"]:
                items = cognitive_map.map.get(d, [])
                dim_summary[d] = len(items)
            cognitive_export = {
                "dimensions": dim_summary,
                "patterns": [p.get("insight", "") for p in cognitive_map.detect_patterns()],
            }
            result = experience_pool.export_crystals(
                rules_to_export, export_file,
                hexagram_data=hexagram_export,
                cognitive_data=cognitive_export,
                contributor_id=contribution.get_contributor_id())
            if result:
                print(f"\n已导出{len(rules_to_export)}条经验 → {result}")
                print("把这个文件发给其他Agent，他用 import 命令导入")
                contribution.add_points("export")
                ecosystem.record_action("export")
            else:
                print("\n导出失败")
            continue

        # 区分 Python "import xxx" 和导入命令 "import 文件路径/.taiji"
        _import_arg = user_input.split(maxsplit=1)[1].strip().strip('"') if len(user_input.split(maxsplit=1)) > 1 else ""
        _is_import_cmd = (user_input.lower() == "import" or
                          (user_input.lower().startswith("import ") and
                           ("/" in _import_arg or "\\" in _import_arg or
                            _import_arg.endswith((".taiji", ".json")))))
        if _is_import_cmd:
            parts = user_input.split(maxsplit=1)
            if len(parts) < 2:
                print("\n用法：import 文件路径")
                print("把朋友发给你的 .taiji 文件拖进来")
                try:
                    imp_path = input("拖入文件 → ").strip().strip('"')
                except EOFError:
                    continue
            else:
                imp_path = parts[1].strip().strip('"')
            if imp_path and Path(imp_path).exists():
                count = experience_pool.import_crystals(imp_path)
                if count > 0:
                    print(f"\n导入成功！新增{count}条共享经验")
                    print("  来自其他Agent的经验已融入你的认知系统")
                    # 检查是否有Agent快照（v2格式）
                    snaps = experience_pool.get_agent_snapshots()
                    if snaps:
                        latest_id = list(snaps.keys())[-1]
                        latest = snaps[latest_id]
                        if latest.get("hexagram", {}).get("current"):
                            hex_name = latest["hexagram"]["current"]
                            print(f"  该Agent当前卦象：{hex_name}")
                        if latest.get("soul", {}).get("patterns"):
                            print(f"  该Agent认知洞察：{latest['soul']['patterns'][0][:40]}...")
                        # 记录到生态网络
                        ecosystem.record_peer(latest_id, {"rules_count": count})
                    contribution.add_points("import")
                    ecosystem.record_action("import")
                    system = rebuild_system()
                else:
                    print("\n没有新经验（可能已经导入过了）")
            else:
                print("\n文件不存在")
            continue

        # 记录上一轮outcome（用当前输入推断上一轮质量）
        if prev_user_input and prev_reply:
            learner.record_outcome(prev_user_input, prev_reply, user_input)

            # 检查是否该结晶
            if learner.should_crystallize():
                # 检查结晶数量限制
                current_count = len(crystallizer.get_active_rules())
                allowed, limit_msg = premium.check_crystal_limit(current_count)
                if allowed:
                    new_crystals = crystallizer.crystallize()
                    if new_crystals:
                        total = len(crystallizer.get_active_rules())
                        shared_count = len(experience_pool.get_shared_rules())
                        print(f"\n  ━━ 经验结晶 ━━")
                        for c in new_crystals:
                            print(f"    ✦ {c['rule']}")
                        if total <= 3:
                            print(f"\n  💡 军师从对话中发现了你的独特规律。")
                            print("    每个人的结晶都不一样——这是专属于你的认知进化。")
                            print("    输入 export 可分享给其他Agent，集体学习。")
                        elif total <= 8:
                            print(f"\n  已积累{total}条结晶。")
                            if shared_count > 0:
                                print(f"    你也在受益于{shared_count}条他人共享的经验。")
                            else:
                                print("    输入 export 导出经验包，让更多人受益。")
                        else:
                            print(f"\n  {total}条结晶在线。")
                            print("    输入 export 导出你的经验包，")
                            print("    别人 import 后你们的经验会交叉验证，共同进化。")
                        contribution.add_points("crystal", len(new_crystals))
                        ecosystem.record_action("crystal", len(new_crystals))
                elif limit_msg:
                    print(f"\n  {limit_msg}")

        # 收集最近用户消息用于卦象更新
        recent_user_msgs = [
            m["content"] for m in history if m["role"] == "user"
        ]
        recent_user_msgs.append(user_input)

        # 更新卦象（从对话状态诊断）
        positive_rate = learner.get_positive_rate()
        hex_result = hexagram_engine.update_from_conversation(
            recent_user_msgs, positive_rate)

        # 每3轮触发一次易经推演（Premium功能）
        round_count = len(history) // 2 + 1
        if round_count >= 3 and round_count % 3 == 0:
            divination = hexagram_engine.divine(recent_user_msgs, positive_rate)
            if divination and divination.get("display"):
                if premium.limits["hex_trend"]:
                    print(divination["display"])
                elif round_count == 3:
                    print("  [卦象趋势] 升级Premium查看完整走势推演 → 输入 upgrade")

        # 更新认知地图（从当前对话提取）
        # 先用空reply，等AI回复后再提取完整的
        cognitive_map.extract_from_message(user_input, "")

        # 意图检测 → 主动执行模式
        intent_prompt = detect_intent(user_input)
        if intent_prompt:
            intent_tag = intent_prompt.split("：")[1].split("】")[0] if "：" in intent_prompt else "分析"
            print(f"\n  [触发] 军师进入{intent_tag}...")

        # 知识库检索
        knowledge_prompt = knowledge.get_knowledge_prompt(user_input)
        if knowledge_prompt:
            print(f"  [知识库] 找到相关参考资料")

        # 重建system prompt（每轮更新，注入最新卦象+认知+意图+知识）
        system = rebuild_system(intent_prompt, knowledge_prompt)

        print("\nAI：", end="", flush=True)
        try:
            reply = chat(system, history, user_input, model_config)
        except Exception as e:
            err = str(e)
            if "401" in err or "authentication" in err.lower() or "api key" in err.lower():
                print(f"\n[错误] API Key无效或已过期")
                print(f"  输入 model 重新配置，或检查你的Key是否正确")
            elif "429" in err or "rate" in err.lower() or "quota" in err.lower():
                print(f"\n[错误] 请求太频繁或额度用完了")
                print(f"  等几秒再试，或去充值/换个模型（输入 model）")
            elif "timeout" in err.lower() or "connect" in err.lower():
                print(f"\n[错误] 网络连接失败")
                print(f"  检查网络，或换个模型试试（输入 model）")
            else:
                print(f"\n[错误] {e}")
            continue

        print(reply)

        # AI回复后，再次更新认知地图（带完整reply）
        cognitive_map.extract_from_message(user_input, reply)

        # 对话积分
        contribution.add_points("chat")
        ecosystem.record_action("chat")

        # 检查成就解锁
        new_achievements = ecosystem.check_achievements(ecosystem.get_stats())
        for a in new_achievements:
            print(f"\n  ★ 成就解锁：{a['name']} — {a['desc']}（+{a['points']}分）")
            contribution.add_points("chat", a["points"])

        # 里程碑引导（集体学习版）
        crystal_count = len(crystallizer.get_active_rules())
        shared_count = len(experience_pool.get_shared_rules())
        if round_count == 5:
            print("\n  ━━ 引擎就绪 ━━")
            print("  5轮对话完成，进化引擎已启动。")
            print("  军师正在学习你的思维模式，形成专属经验。")
            if shared_count > 0:
                print(f"  {shared_count}条其他Agent的经验已在为你服务。")
        elif round_count == 10 and crystal_count == 0:
            print("\n  ━━ 即将结晶 ━━")
            print("  10轮了，经验结晶即将生成——")
            print("  每个人的结晶都不一样，这是你独有的认知进化。")
        elif round_count == 15 and crystal_count > 0:
            print(f"\n  ━━ 集体学习 ━━")
            print(f"  你已有{crystal_count}条独有经验。")
            print(f"  输入 export 导出经验包，发给朋友 import，")
            print(f"  互相验证的经验会越来越准——这就是集体进化。")
        elif round_count == 20 and crystal_count >= 3:
            print(f"\n  ━━ 经验网络 ━━")
            print(f"  {crystal_count}条结晶在线。每条都是你独有的规律。")
            print(f"  导出给其他Agent后，经验会被交叉验证，")
            print(f"  多人验证过的经验=集体智慧，比单人更可靠。")

        history.append({"role": "user",      "content": user_input})
        history.append({"role": "assistant", "content": reply})

        max_history = premium.limits["max_history"]
        if len(history) > max_history:
            # 压缩前先提取老对话的要点，注入认知地图
            old_msgs = history[:len(history) - max_history]
            old_user_msgs = [m["content"] for m in old_msgs if m["role"] == "user"]
            for om in old_user_msgs:
                cognitive_map.extract_from_message(om, "")
            history = history[-max_history:]

        save_history(history_key, history)

        prev_user_input = user_input
        prev_reply = reply

    # 退出前记录最后一轮
    if prev_user_input and prev_reply:
        learner.record_outcome(prev_user_input, prev_reply)

    # 退出前结晶
    new_crystals = crystallizer.crystallize()
    if new_crystals:
        print(f"\n[进化] 退出时新增{len(new_crystals)}条经验结晶")

    stats = learner.get_stats_display()
    if stats:
        print(f"\n{stats}")

    print("\n对话已保存 | 经验已结晶")
    try:
        input("\n按回车退出...")
    except EOFError:
        pass

if __name__ == "__main__":
    if sys.platform == "win32":
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
        # Windows 终端粘贴多行文本时可能触发编码异常，设置 stdin 也用 utf-8
        try:
            sys.stdin = io.TextIOWrapper(sys.stdin.buffer, encoding="utf-8", errors="replace")
        except Exception:
            pass
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n退出")
    except Exception as e:
        print(f"\n\n程序遇到错误：{e}")
        print("你的对话历史已自动保存，下次启动会恢复。")
        try:
            input("\n按回车退出...")
        except Exception:
            pass
